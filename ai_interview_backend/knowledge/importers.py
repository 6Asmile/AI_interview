import csv
import io
import os
import tempfile
from dataclasses import dataclass, field
from importlib.metadata import PackageNotFoundError, version

import openpyxl
from docx import Document
from pypdf import PdfReader

from .models import KnowledgeDocument


SUPPORTED_EXTENSIONS = {'.md', '.txt', '.pdf', '.docx', '.xlsx', '.csv', '.png', '.jpg', '.jpeg', '.webp', '.bmp'}
IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.webp', '.bmp'}
FAQ_QUESTION_COLUMNS = {'question', '问题', '题目', 'q'}
FAQ_ANSWER_COLUMNS = {'answer', '答案', '回答', 'a'}
JOB_COLUMNS = {'job_position', 'job_positions', '岗位', '适用岗位'}
TAG_COLUMNS = {'ability_tags', 'tags', '能力标签', '标签'}
DIFFICULTY_COLUMNS = {'difficulty', '难度'}


@dataclass
class ParsedKnowledgeBlock:
    block_type: str
    text: str
    heading_path: list[str] = field(default_factory=list)
    page_start: int | None = None
    page_end: int | None = None
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedKnowledgeDocument:
    title: str
    content: str
    source_type: str = 'question_bank'
    job_positions: list[str] = field(default_factory=list)
    ability_tags: list[str] = field(default_factory=list)
    difficulty: str = KnowledgeDocument.Difficulty.ANY
    file_type: str = ''
    parsed_content: dict = field(default_factory=dict)
    parser_name: str = ''
    parser_version: str = ''
    parser_fallback_reason: str = ''
    ocr_enabled: bool = False


def _package_version(package_name: str) -> str:
    try:
        return version(package_name)
    except PackageNotFoundError:
        return ''


def _read_bytes(uploaded_file) -> bytes:
    uploaded_file.seek(0)
    data = uploaded_file.read()
    uploaded_file.seek(0)
    return data


def _decode_text(data: bytes) -> str:
    for encoding in ('utf-8-sig', 'utf-8', 'gb18030'):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode('utf-8', errors='ignore')


def _split_list(value) -> list[str]:
    if value is None:
        return []
    if isinstance(value, (list, tuple)):
        items = value
    else:
        text = str(value)
        for separator in ['；', ';', '|', '，']:
            text = text.replace(separator, ',')
        items = text.split(',')
    return [str(item).strip() for item in items if str(item).strip()]


def _safe_difficulty(value: str) -> str:
    value = (value or '').strip().lower()
    labels = {
        '不限': KnowledgeDocument.Difficulty.ANY,
        '基础': KnowledgeDocument.Difficulty.EASY,
        '中等': KnowledgeDocument.Difficulty.MEDIUM,
        '高阶': KnowledgeDocument.Difficulty.HARD,
    }
    value = labels.get(value, value)
    valid = {choice[0] for choice in KnowledgeDocument.Difficulty.choices}
    return value if value in valid else KnowledgeDocument.Difficulty.ANY


def _normalized_header_map(headers: list[str]) -> dict[str, int]:
    return {str(header or '').strip().lower(): index for index, header in enumerate(headers)}


def _find_column(header_map: dict[str, int], candidates: set[str]) -> int | None:
    normalized = {item.lower() for item in candidates}
    for header, index in header_map.items():
        if header in normalized:
            return index
    return None


def _blocks_to_content(blocks: list[ParsedKnowledgeBlock]) -> str:
    return '\n\n'.join(block.text.strip() for block in blocks if block.text.strip())


def _serialize_blocks(blocks: list[ParsedKnowledgeBlock], parser_name: str, parser_version: str, fallback_reason: str = '') -> dict:
    return {
        'parser_name': parser_name,
        'parser_version': parser_version,
        'parser_fallback_reason': fallback_reason,
        'blocks': [
            {
                'block_type': block.block_type,
                'text': block.text,
                'heading_path': block.heading_path,
                'page_start': block.page_start,
                'page_end': block.page_end,
                'metadata': block.metadata,
            }
            for block in blocks
            if block.text.strip()
        ],
    }


def _save_temp_file(uploaded_file, suffix: str) -> str:
    data = _read_bytes(uploaded_file)
    handle = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    try:
        handle.write(data)
        return handle.name
    finally:
        handle.close()


def _run_paddleocr(path: str, lang: str = 'ch') -> tuple[str, list[dict]]:
    try:
        from paddleocr import PaddleOCR
    except Exception as exc:
        raise RuntimeError(f'PaddleOCR 不可用：{exc}') from exc

    ocr = PaddleOCR(use_angle_cls=True, lang=lang)
    raw_result = ocr.ocr(path, cls=True)
    lines = []
    items = []
    for page in raw_result or []:
        for row in page or []:
            if not row or len(row) < 2:
                continue
            text = row[1][0] if row[1] else ''
            confidence = float(row[1][1]) if row[1] and len(row[1]) > 1 else None
            if text:
                lines.append(text)
                items.append({'text': text, 'confidence': confidence, 'bbox': row[0]})
    return '\n'.join(lines), items


class DocumentParsingService:
    def __init__(self, *, enable_ocr: bool = True, ocr_lang: str = 'ch'):
        self.enable_ocr = enable_ocr
        self.ocr_lang = ocr_lang

    def parse(self, uploaded_file) -> ParsedKnowledgeDocument:
        name = getattr(uploaded_file, 'name', '') or '未命名文件'
        title, extension = os.path.splitext(os.path.basename(name))
        extension = extension.lower()
        if extension not in SUPPORTED_EXTENSIONS:
            raise ValueError(f'不支持的文件类型：{extension or "unknown"}')

        docling_error = ''
        if extension in {'.pdf', '.docx', '.xlsx'}:
            try:
                return self._parse_with_docling(uploaded_file, title or name, extension)
            except Exception as exc:
                docling_error = f'Docling 解析失败，已降级：{exc}'

        if extension in IMAGE_EXTENSIONS:
            if not self.enable_ocr:
                raise ValueError('图片文件需要 OCR，但 OCR 未启用。')
            return self._parse_image_with_ocr(uploaded_file, title or name, extension, docling_error)
        if extension in {'.md', '.txt'}:
            return self._parse_text(uploaded_file, title or name, extension, docling_error)
        if extension == '.pdf':
            return self._parse_pdf_fallback(uploaded_file, title or name, docling_error)
        if extension == '.docx':
            return self._parse_docx_fallback(uploaded_file, title or name, docling_error)
        if extension == '.xlsx':
            return self._parse_xlsx_fallback(uploaded_file, title or name, docling_error)
        if extension == '.csv':
            return self._parse_csv(uploaded_file, title or name, docling_error)
        raise ValueError(f'不支持的文件类型：{extension}')

    def _parse_with_docling(self, uploaded_file, title: str, extension: str) -> ParsedKnowledgeDocument:
        try:
            from docling.document_converter import DocumentConverter
        except Exception as exc:
            raise RuntimeError(f'Docling 未安装或不可导入：{exc}') from exc

        path = _save_temp_file(uploaded_file, extension)
        try:
            converter = DocumentConverter()
            result = converter.convert(path)
            document = result.document
            markdown = document.export_to_markdown()
            blocks = self._blocks_from_markdown(markdown, default_type='paragraph')
            ocr_text = ''
            ocr_items = []
            if self.enable_ocr and extension in IMAGE_EXTENSIONS:
                try:
                    ocr_text, ocr_items = _run_paddleocr(path, lang=self.ocr_lang)
                except Exception as exc:
                    blocks.append(ParsedKnowledgeBlock(
                        block_type='ocr_error',
                        text='',
                        metadata={'ocr_engine': 'paddleocr', 'error': str(exc)},
                    ))
            if ocr_text:
                blocks.append(ParsedKnowledgeBlock(
                    block_type='ocr',
                    text=ocr_text,
                    metadata={'ocr_engine': 'paddleocr', 'ocr_items': ocr_items},
                ))
            content = _blocks_to_content(blocks)
            if not content:
                raise ValueError('Docling 未提取到文本内容。')
            parser_version = _package_version('docling')
            return ParsedKnowledgeDocument(
                title=title,
                content=content,
                source_type='document',
                file_type=extension.lstrip('.'),
                parsed_content=_serialize_blocks(blocks, 'docling', parser_version),
                parser_name='docling',
                parser_version=parser_version,
                ocr_enabled=bool(ocr_text),
            )
        finally:
            try:
                os.unlink(path)
            except OSError:
                pass

    def _blocks_from_markdown(self, markdown: str, default_type: str = 'paragraph') -> list[ParsedKnowledgeBlock]:
        blocks = []
        heading_path: list[str] = []
        buffer: list[str] = []

        def flush_buffer():
            if buffer:
                blocks.append(ParsedKnowledgeBlock(
                    block_type=default_type,
                    text='\n'.join(buffer).strip(),
                    heading_path=list(heading_path),
                ))
                buffer.clear()

        for line in (markdown or '').splitlines():
            stripped = line.strip()
            if not stripped:
                flush_buffer()
                continue
            if stripped.startswith('#'):
                flush_buffer()
                level = len(stripped) - len(stripped.lstrip('#'))
                heading = stripped[level:].strip()
                heading_path[:] = heading_path[:max(level - 1, 0)]
                heading_path.append(heading)
                blocks.append(ParsedKnowledgeBlock('heading', heading, list(heading_path)))
            elif stripped.startswith('|') and stripped.endswith('|'):
                flush_buffer()
                blocks.append(ParsedKnowledgeBlock('table', stripped, list(heading_path)))
            else:
                buffer.append(stripped)
        flush_buffer()
        return blocks

    def _parse_text(self, uploaded_file, title: str, extension: str, fallback_reason: str = '') -> ParsedKnowledgeDocument:
        content = _decode_text(_read_bytes(uploaded_file)).strip()
        if not content:
            raise ValueError('文本文件为空。')
        blocks = self._blocks_from_markdown(content) if extension == '.md' else [
            ParsedKnowledgeBlock('paragraph', content)
        ]
        parser_name = 'markdown' if extension == '.md' else 'plain_text'
        return ParsedKnowledgeDocument(
            title=title,
            content=_blocks_to_content(blocks),
            source_type='document' if extension == '.md' else 'question_bank',
            file_type=extension.lstrip('.'),
            parsed_content=_serialize_blocks(blocks, parser_name, '', fallback_reason),
            parser_name=parser_name,
            parser_fallback_reason=fallback_reason,
        )

    def _parse_image_with_ocr(self, uploaded_file, title: str, extension: str, fallback_reason: str = '') -> ParsedKnowledgeDocument:
        path = _save_temp_file(uploaded_file, extension)
        try:
            text, items = _run_paddleocr(path, lang=self.ocr_lang)
        finally:
            try:
                os.unlink(path)
            except OSError:
                pass
        if not text.strip():
            raise ValueError('图片 OCR 未识别到文本。')
        blocks = [ParsedKnowledgeBlock(
            block_type='ocr',
            text=text,
            metadata={'ocr_engine': 'paddleocr', 'ocr_items': items},
        )]
        return ParsedKnowledgeDocument(
            title=title,
            content=text,
            source_type='document',
            file_type=extension.lstrip('.'),
            parsed_content=_serialize_blocks(blocks, 'paddleocr', _package_version('paddleocr'), fallback_reason),
            parser_name='paddleocr',
            parser_version=_package_version('paddleocr'),
            parser_fallback_reason=fallback_reason,
            ocr_enabled=True,
        )

    def _parse_pdf_fallback(self, uploaded_file, title: str, fallback_reason: str = '') -> ParsedKnowledgeDocument:
        uploaded_file.seek(0)
        reader = PdfReader(uploaded_file)
        pages = [page.extract_text() or '' for page in reader.pages]
        uploaded_file.seek(0)
        blocks = [
            ParsedKnowledgeBlock('page', page.strip(), page_start=index, page_end=index)
            for index, page in enumerate(pages, start=1)
            if page.strip()
        ]
        content = _blocks_to_content(blocks)
        if not content:
            raise ValueError('PDF 未提取到文本；扫描版 PDF 需要 OCR。')
        return ParsedKnowledgeDocument(
            title=title,
            content=content,
            source_type='document',
            file_type='pdf',
            parsed_content=_serialize_blocks(blocks, 'pypdf', _package_version('pypdf'), fallback_reason),
            parser_name='pypdf',
            parser_version=_package_version('pypdf'),
            parser_fallback_reason=fallback_reason,
        )

    def _parse_docx_fallback(self, uploaded_file, title: str, fallback_reason: str = '') -> ParsedKnowledgeDocument:
        uploaded_file.seek(0)
        document = Document(uploaded_file)
        uploaded_file.seek(0)
        blocks = [
            ParsedKnowledgeBlock('paragraph', paragraph.text.strip())
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        for table in document.tables:
            rows = []
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    rows.append(' | '.join(values))
            if rows:
                blocks.append(ParsedKnowledgeBlock('table', '\n'.join(rows)))
        content = _blocks_to_content(blocks)
        if not content:
            raise ValueError('Word 文件没有可导入的文本内容。')
        return ParsedKnowledgeDocument(
            title=title,
            content=content,
            source_type='document',
            file_type='docx',
            parsed_content=_serialize_blocks(blocks, 'python-docx', _package_version('python-docx'), fallback_reason),
            parser_name='python-docx',
            parser_version=_package_version('python-docx'),
            parser_fallback_reason=fallback_reason,
        )

    def _parse_csv(self, uploaded_file, title: str, fallback_reason: str = '') -> ParsedKnowledgeDocument:
        text = _decode_text(_read_bytes(uploaded_file))
        reader = csv.reader(io.StringIO(text))
        rows = [row for row in reader]
        if not rows:
            raise ValueError('CSV 文件为空。')
        return self._document_from_rows(title, rows, 'csv', fallback_reason)

    def _parse_xlsx_fallback(self, uploaded_file, title: str, fallback_reason: str = '') -> ParsedKnowledgeDocument:
        uploaded_file.seek(0)
        workbook = openpyxl.load_workbook(uploaded_file, data_only=True, read_only=True)
        sheet = workbook.active
        rows = [list(row) for row in sheet.iter_rows(values_only=True)]
        uploaded_file.seek(0)
        if not rows:
            raise ValueError('Excel 文件为空。')
        return self._document_from_rows(title, rows, 'xlsx', fallback_reason)

    def _document_from_rows(self, title: str, rows: list[list], file_type: str, fallback_reason: str = '') -> ParsedKnowledgeDocument:
        headers, body = rows[0], rows[1:]
        header_map = _normalized_header_map(headers)
        question_index = _find_column(header_map, FAQ_QUESTION_COLUMNS)
        answer_index = _find_column(header_map, FAQ_ANSWER_COLUMNS)
        job_index = _find_column(header_map, JOB_COLUMNS)
        tag_index = _find_column(header_map, TAG_COLUMNS)
        difficulty_index = _find_column(header_map, DIFFICULTY_COLUMNS)
        blocks = []
        jobs = []
        tags = []
        difficulty = KnowledgeDocument.Difficulty.ANY
        source_type = 'document'
        if question_index is not None and answer_index is not None:
            source_type = 'faq'
            for number, row in enumerate(body, start=1):
                question = str(row[question_index] or '').strip() if question_index < len(row) else ''
                answer = str(row[answer_index] or '').strip() if answer_index < len(row) else ''
                if not question and not answer:
                    continue
                blocks.append(ParsedKnowledgeBlock(
                    'faq',
                    f'Q: {question}\nA: {answer}',
                    heading_path=[f'FAQ {number}'],
                    metadata={'question': question, 'answer': answer},
                ))
                if job_index is not None and job_index < len(row):
                    jobs.extend(_split_list(row[job_index]))
                if tag_index is not None and tag_index < len(row):
                    tags.extend(_split_list(row[tag_index]))
                if difficulty_index is not None and difficulty_index < len(row) and row[difficulty_index]:
                    difficulty = _safe_difficulty(str(row[difficulty_index]))
        else:
            for row in rows:
                values = [str(value).strip() for value in row if value not in (None, '')]
                if values:
                    blocks.append(ParsedKnowledgeBlock('table', ' | '.join(values)))
        content = _blocks_to_content(blocks)
        if not content:
            raise ValueError('表格文件没有可导入的文本内容。')
        parser_name = 'openpyxl' if file_type == 'xlsx' else 'csv'
        return ParsedKnowledgeDocument(
            title=title,
            content=content,
            source_type=source_type,
            job_positions=sorted(set(jobs)),
            ability_tags=sorted(set(tags)),
            difficulty=difficulty,
            file_type=file_type,
            parsed_content=_serialize_blocks(blocks, parser_name, _package_version(parser_name), fallback_reason),
            parser_name=parser_name,
            parser_version=_package_version(parser_name),
            parser_fallback_reason=fallback_reason,
        )


def parse_knowledge_file(uploaded_file, *, enable_ocr: bool = True, ocr_lang: str = 'ch') -> ParsedKnowledgeDocument:
    return DocumentParsingService(enable_ocr=enable_ocr, ocr_lang=ocr_lang).parse(uploaded_file)
