from __future__ import annotations

import io
import json
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

import yaml
from django.core.files.base import ContentFile
from django.utils import timezone
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt
from PIL import Image

from .models import ResumeArtifact, ResumeAsset
from .runtime import resume_runtime_config
from .schema import canonical_json, sha256_json, strip_internal_metadata


RENDERER_NAME = 'rendercv-typst'
RENDERER_VERSION = '2.8'
UNSAFE_TEXT = re.compile(
    r'(?i)(#[a-z]|!\s*\[|(?:file|javascript|data):|\\\\(?:input|include)|<\s*/?\s*[a-z])'
)
THEME_MAP = {
    'ats-classic': 'classic',
    'modern-professional': 'moderncv',
    'engineering': 'engineeringresumes',
    'graduate': 'engineeringclassic',
    'management-consulting': 'harvard',
    'academic-research': 'sb2nov',
}


class RenderFailure(RuntimeError):
    def __init__(self, code: str, message: str):
        super().__init__(message)
        self.code = code


def artifact_cache_key(
    content_hash: str,
    design_hash: str,
    output_format: str,
    *,
    namespace: str = '',
) -> str:
    return sha256_json({
        'namespace': namespace,
        'content_hash': content_hash,
        'design_hash': design_hash,
        'format': output_format,
        'renderer': RENDERER_NAME,
        'renderer_version': RENDERER_VERSION,
    })


def _safe_text(value) -> str:
    text = str(value or '')
    if len(text) > 20_000:
        raise RenderFailure('unsafe_input', '单个文本字段超过 20,000 字符。')
    if UNSAFE_TEXT.search(text):
        raise RenderFailure('unsafe_input', '文本包含不允许的排版或文件访问指令。')
    return text


def _rendercv_payload(resume_json: dict, design_json: dict, avatar_filename: str = '') -> dict:
    basics = resume_json.get('basics') or {}
    location = basics.get('location') or {}
    section_names = {
        'zh-CN': {
            'summary': '职业摘要', 'work': '工作经历', 'projects': '项目经历',
            'education': '教育经历', 'skills': '专业技能', 'certificates': '证书',
            'awards': '荣誉奖项', 'publications': '发表成果', 'languages': '语言能力',
            'volunteer': '志愿经历', 'interests': '兴趣', 'references': '推荐人',
        },
        'en-US': {
            'summary': 'Summary', 'work': 'Experience', 'projects': 'Projects',
            'education': 'Education', 'skills': 'Skills', 'certificates': 'Certificates',
            'awards': 'Awards', 'publications': 'Publications', 'languages': 'Languages',
            'volunteer': 'Volunteer', 'interests': 'Interests', 'references': 'References',
        },
    }[design_json.get('language', 'zh-CN')]
    sections = {}
    hidden_sections = set(design_json.get('hidden_sections') or [])
    summary = '' if 'summary' in hidden_sections else _safe_text(basics.get('summary'))
    if summary:
        sections[section_names['summary']] = [summary]
    work_entries = []
    for item in resume_json.get('work') or []:
        work_entries.append({
            'company': _safe_text(item.get('name')),
            'position': _safe_text(item.get('position')),
            'location': _safe_text(item.get('location')),
            'start_date': _safe_text(item.get('startDate')),
            'end_date': _safe_text(item.get('endDate')) or 'present',
            'summary': _safe_text(item.get('summary')),
            'highlights': [_safe_text(value) for value in item.get('highlights') or []],
        })
    if work_entries and 'work' not in hidden_sections:
        sections[section_names['work']] = work_entries
    project_entries = []
    for item in resume_json.get('projects') or []:
        details = ', '.join(_safe_text(value) for value in item.get('keywords') or [])
        highlights = [_safe_text(value) for value in item.get('highlights') or []]
        description = _safe_text(item.get('description'))
        if description:
            highlights.insert(0, description)
        project_entry = {
            'name': _safe_text(item.get('name')),
            'location': details,
            'highlights': highlights,
        }
        if item.get('startDate'):
            project_entry['start_date'] = _safe_text(item.get('startDate'))
        if item.get('endDate'):
            project_entry['end_date'] = _safe_text(item.get('endDate'))
        project_entries.append(project_entry)
    if project_entries and 'projects' not in hidden_sections:
        sections[section_names['projects']] = project_entries
    education_entries = []
    for item in resume_json.get('education') or []:
        education_entry = {
            'institution': _safe_text(item.get('institution')),
            'area': _safe_text(item.get('area')),
            'degree': _safe_text(item.get('studyType')),
            'highlights': [_safe_text(value) for value in item.get('courses') or []],
        }
        if item.get('startDate'):
            education_entry['start_date'] = _safe_text(item.get('startDate'))
        if item.get('endDate'):
            education_entry['end_date'] = _safe_text(item.get('endDate'))
        education_entries.append(education_entry)
    if education_entries and 'education' not in hidden_sections:
        sections[section_names['education']] = education_entries
    skills = []
    for item in resume_json.get('skills') or []:
        details = ', '.join(_safe_text(value) for value in item.get('keywords') or [])
        if item.get('level'):
            details = ', '.join(filter(None, [_safe_text(item.get('level')), details]))
        skills.append({'label': _safe_text(item.get('name')), 'details': details})
    if skills and 'skills' not in hidden_sections:
        sections[section_names['skills']] = skills
    for key in ('certificates', 'awards', 'publications', 'languages', 'volunteer', 'interests'):
        if key in hidden_sections:
            continue
        entries = []
        for item in resume_json.get(key) or []:
            values = [
                _safe_text(value)
                for item_key, value in item.items()
                if item_key != 'x-ifaceoff' and isinstance(value, (str, int, float)) and value
            ]
            if values:
                entries.append(' — '.join(values))
        if entries:
            sections[section_names[key]] = entries
    if 'references' not in hidden_sections:
        entries = []
        for item in resume_json.get('references') or []:
            text = ' — '.join(filter(None, [
                _safe_text(item.get('name')),
                _safe_text(item.get('reference')),
            ]))
            if text:
                entries.append(text)
        if entries:
            sections[section_names['references']] = entries
    ordered = {}
    for section_key in design_json.get('section_order') or []:
        translated = section_names.get(section_key)
        if translated and translated in sections:
            ordered[translated] = sections[translated]
    for title, entries in sections.items():
        ordered.setdefault(title, entries)
    location_text = ', '.join(filter(None, [
        _safe_text(location.get('city')),
        _safe_text(location.get('region')),
        _safe_text(location.get('countryCode')),
    ]))
    density = design_json.get('density', 'balanced')
    density_options = {
        'compact': {'margin': '0.45in', 'body_size': '9pt', 'line_spacing': '0.45em'},
        'balanced': {'margin': '0.6in', 'body_size': '10pt', 'line_spacing': '0.6em'},
        'comfortable': {'margin': '0.75in', 'body_size': '10.5pt', 'line_spacing': '0.72em'},
    }[density if density in {'compact', 'balanced', 'comfortable'} else 'balanced']
    color = design_json.get('color', '#1F2937')
    page_size = 'us-letter' if design_json.get('page_size') == 'Letter' else 'a4'
    date_template = {
        'YYYY-MM': 'YEAR-MONTH_NUMBER',
        'YYYY.MM': 'YEAR.MONTH_NUMBER',
        'MMM YYYY': 'MONTH_ABBREVIATION YEAR',
        'YYYY': 'YEAR',
    }.get(design_json.get('date_format'), 'YEAR-MONTH_NUMBER')
    cv = {
            'name': _safe_text(basics.get('name')) or 'Resume',
            'headline': _safe_text(basics.get('label')),
            'location': location_text,
            'email': _safe_text(basics.get('email')) or None,
            'phone': _safe_text(basics.get('phone')) or None,
            'website': _safe_text(basics.get('url')) or None,
            'sections': ordered,
        }
    if avatar_filename and design_json.get('show_avatar'):
        cv['photo'] = avatar_filename
    return {
        'cv': cv,
        'design': {
            'theme': THEME_MAP.get(design_json.get('template_key'), 'classic'),
            'page': {
                'size': page_size,
                'top_margin': density_options['margin'],
                'bottom_margin': density_options['margin'],
                'left_margin': density_options['margin'],
                'right_margin': density_options['margin'],
                'show_top_note': False,
            },
            'colors': {
                'name': color,
                'headline': color,
                'connections': color,
                'section_titles': color,
                'links': color,
            },
            'typography': {
                'font_family': design_json.get('font', 'Source Sans 3'),
                'font_size': {'body': density_options['body_size']},
                'line_spacing': density_options['line_spacing'],
            },
            'templates': {
                'single_date': date_template,
                'date_range': f'START_DATE – END_DATE',
            },
        },
        'locale': {'language': 'mandarin_chinese' if design_json.get('language') == 'zh-CN' else 'english'},
    }


def _docx_bytes(resume_json: dict) -> bytes:
    data = strip_internal_metadata(resume_json)
    basics = data.get('basics') or {}
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    normal = document.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(10)
    title = document.add_heading(_safe_text(basics.get('name')) or 'Resume', level=0)
    title.paragraph_format.space_after = Pt(2)
    contact = ' | '.join(filter(None, [
        _safe_text(basics.get('label')), _safe_text(basics.get('email')),
        _safe_text(basics.get('phone')), _safe_text(basics.get('url')),
    ]))
    if contact:
        document.add_paragraph(contact)
    if basics.get('summary'):
        document.add_heading('Summary', level=1)
        document.add_paragraph(_safe_text(basics['summary']))
    mappings = [
        ('work', 'Experience', 'name', 'position'),
        ('projects', 'Projects', 'name', 'description'),
        ('education', 'Education', 'institution', 'area'),
        ('skills', 'Skills', 'name', 'level'),
        ('certificates', 'Certificates', 'name', 'issuer'),
        ('awards', 'Awards', 'title', 'awarder'),
        ('publications', 'Publications', 'name', 'publisher'),
        ('languages', 'Languages', 'language', 'fluency'),
    ]
    for key, heading, primary, secondary in mappings:
        items = data.get(key) or []
        if not items:
            continue
        document.add_heading(heading, level=1)
        for item in items:
            line = ' — '.join(filter(None, [_safe_text(item.get(primary)), _safe_text(item.get(secondary))]))
            dates = ' – '.join(filter(None, [_safe_text(item.get('startDate')), _safe_text(item.get('endDate'))]))
            paragraph = document.add_paragraph()
            paragraph.add_run(line).bold = True
            if dates:
                paragraph.add_run(f'  {dates}')
            detail = _safe_text(item.get('summary') or item.get('description'))
            if detail:
                document.add_paragraph(detail)
            for highlight in item.get('highlights') or item.get('keywords') or []:
                document.add_paragraph(_safe_text(highlight), style='List Bullet')
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _markdown_bytes(resume_json: dict) -> bytes:
    """Generate deterministic, ATS-friendly Markdown without internal IDs."""
    data = strip_internal_metadata(resume_json)
    basics = data.get('basics') or {}
    lines = [f"# {_safe_text(basics.get('name')) or 'Resume'}", '']
    contact = ' · '.join(filter(None, [
        _safe_text(basics.get('label')), _safe_text(basics.get('email')),
        _safe_text(basics.get('phone')), _safe_text(basics.get('url')),
    ]))
    if contact:
        lines.extend([contact, ''])
    if basics.get('summary'):
        lines.extend(['## 职业摘要', '', _safe_text(basics['summary']), ''])
    headings = {
        'work': '工作经历', 'projects': '项目经历', 'education': '教育经历', 'skills': '专业技能',
        'certificates': '证书', 'awards': '荣誉奖项', 'publications': '发表成果',
        'languages': '语言能力', 'volunteer': '志愿经历', 'interests': '兴趣', 'references': '推荐人',
    }
    primary_fields = {
        'work': ('name', 'position'), 'projects': ('name', 'description'),
        'education': ('institution', 'area'), 'skills': ('name', 'level'),
        'certificates': ('name', 'issuer'), 'awards': ('title', 'awarder'),
        'publications': ('name', 'publisher'), 'languages': ('language', 'fluency'),
        'volunteer': ('organization', 'position'), 'interests': ('name', ''),
        'references': ('name', 'reference'),
    }
    for key, heading in headings.items():
        items = data.get(key) or []
        if not items:
            continue
        lines.extend([f'## {heading}', ''])
        primary, secondary = primary_fields[key]
        for item in items:
            title = ' — '.join(filter(None, [_safe_text(item.get(primary)), _safe_text(item.get(secondary))]))
            dates = ' – '.join(filter(None, [_safe_text(item.get('startDate')), _safe_text(item.get('endDate'))]))
            lines.append(f"### {title or heading}{f' · {dates}' if dates else ''}")
            detail = _safe_text(item.get('summary') or item.get('description'))
            if detail:
                lines.append(detail)
            for bullet in item.get('highlights') or item.get('keywords') or []:
                lines.append(f'- {_safe_text(bullet)}')
            lines.append('')
    return ('\n'.join(lines).rstrip() + '\n').encode('utf-8')


def _stack_png_pages(paths: list[Path]) -> bytes:
    images = [Image.open(path).convert('RGB') for path in paths]
    if not images:
        raise RenderFailure('renderer_output_missing', 'RenderCV 未生成 PNG。')
    width = max(image.width for image in images)
    gap = 24 if len(images) > 1 else 0
    height = sum(image.height for image in images) + gap * (len(images) - 1)
    canvas = Image.new('RGB', (width, height), 'white')
    cursor = 0
    for image in images:
        canvas.paste(image, ((width - image.width) // 2, cursor))
        cursor += image.height + gap
    output = io.BytesIO()
    canvas.save(output, format='PNG', optimize=True)
    return output.getvalue()


def _render_with_rendercv(
    resume_json: dict,
    design_json: dict,
    output_format: str,
    *,
    avatar_bytes: bytes | None = None,
) -> tuple[bytes, str, int]:
    configured_executable = os.environ.get('RESUME_RENDER_EXECUTABLE', '').strip()
    bundled_executable = Path('/opt/rendercv/bin/rendercv')
    executable = (
        configured_executable
        or (str(bundled_executable) if bundled_executable.exists() else '')
        or shutil.which('rendercv')
    )
    if not executable:
        raise RenderFailure('renderer_unavailable', 'RenderCV 2.8 未安装或不在 PATH 中。')
    with tempfile.TemporaryDirectory(prefix='ifaceoff-resume-') as temp_dir:
        base = Path(temp_dir)
        yaml_path = base / 'resume_CV.yaml'
        pdf_path = base / 'resume.pdf'
        png_path = base / 'resume.png'
        avatar_filename = ''
        if avatar_bytes and design_json.get('show_avatar'):
            avatar_filename = 'avatar.png'
            (base / avatar_filename).write_bytes(avatar_bytes)
        yaml_path.write_text(
            yaml.safe_dump(
                _rendercv_payload(resume_json, design_json, avatar_filename),
                allow_unicode=True,
                sort_keys=False,
            ),
            encoding='utf-8',
        )
        command = [
            executable, 'render', str(yaml_path),
            '--pdf-path', str(pdf_path),
            '--png-path', str(png_path),
            '--dont-generate-html', '--dont-generate-markdown',
        ]
        environment = {
            'PATH': os.environ.get('PATH', ''),
            'PATHEXT': os.environ.get('PATHEXT', ''),
            'SYSTEMROOT': os.environ.get('SYSTEMROOT', ''),
            'APPDATA': os.environ.get('APPDATA', ''),
            'LOCALAPPDATA': os.environ.get('LOCALAPPDATA', ''),
            'TEMP': temp_dir,
            'TMP': temp_dir,
            'TMPDIR': temp_dir,
            'NO_PROXY': '*',
            'HTTP_PROXY': '',
            'HTTPS_PROXY': '',
        }
        if os.name == 'nt':
            environment.update({
                'USERPROFILE': os.environ.get('USERPROFILE', ''),
                'HOMEDRIVE': os.environ.get('HOMEDRIVE', ''),
                'HOMEPATH': os.environ.get('HOMEPATH', ''),
            })
        else:
            environment['HOME'] = temp_dir
        render_cache_dir = (
            os.environ.get('RESUME_RENDER_CACHE_DIR', '').strip()
            or os.environ.get('XDG_CACHE_HOME', '').strip()
            or os.environ.get('LOCALAPPDATA', '').strip()
        )
        if render_cache_dir:
            environment['XDG_CACHE_HOME'] = render_cache_dir
        try:
            completed = subprocess.run(
                command,
                cwd=temp_dir,
                env=environment,
                capture_output=True,
                text=True,
                timeout=int(resume_runtime_config().get('render_timeout_seconds', 20)),
                check=False,
            )
        except subprocess.TimeoutExpired as exc:
            raise RenderFailure('renderer_timeout', '简历渲染超过配置的时间限制。') from exc
        if completed.returncode != 0:
            message = (completed.stderr or completed.stdout or 'RenderCV failed')[-2000:]
            raise RenderFailure('renderer_failed', message)
        if output_format == ResumeArtifact.Format.PDF:
            if not pdf_path.exists():
                raise RenderFailure('renderer_output_missing', 'RenderCV 未生成 PDF。')
            return pdf_path.read_bytes(), 'application/pdf', len(list(base.glob('resume_*.png'))) or 1
        candidates = sorted(base.glob('resume_*.png')) or ([png_path] if png_path.exists() else [])
        if not candidates:
            raise RenderFailure('renderer_output_missing', 'RenderCV 未生成预览图。')
        if output_format == ResumeArtifact.Format.PNG:
            return _stack_png_pages(candidates), 'image/png', len(candidates)
        return candidates[0].read_bytes(), 'image/png', len(candidates)


def render_artifact(artifact: ResumeArtifact) -> ResumeArtifact:
    artifact.status = ResumeArtifact.Status.PROCESSING
    artifact.error_code = ''
    artifact.error_message = ''
    artifact.save(update_fields=['status', 'error_code', 'error_message'])
    resume_json = artifact.preview_input or artifact.content_version.resume_json
    design_json = artifact.preview_design or artifact.design_revision.design_json
    if artifact.format == ResumeArtifact.Format.JSON:
        content = (canonical_json(strip_internal_metadata(resume_json)) + '\n').encode('utf-8')
        mime_type, extension, page_count = 'application/json', 'json', 0
    elif artifact.format == ResumeArtifact.Format.DOCX:
        content = _docx_bytes(resume_json)
        mime_type, extension, page_count = (
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'docx', 0,
        )
    elif artifact.format == ResumeArtifact.Format.MARKDOWN:
        content = _markdown_bytes(resume_json)
        mime_type, extension, page_count = 'text/markdown; charset=utf-8', 'md', 0
    elif artifact.format in {ResumeArtifact.Format.PDF, ResumeArtifact.Format.PNG, ResumeArtifact.Format.PREVIEW}:
        avatar_bytes = None
        image_pointer = str((resume_json.get('basics') or {}).get('image') or '')
        if design_json.get('show_avatar') and image_pointer.startswith('asset:'):
            try:
                asset_id = int(image_pointer.split(':', 1)[1])
            except ValueError:
                asset_id = 0
            avatar = ResumeAsset.objects.filter(
                pk=asset_id,
                resume=artifact.resume,
                kind=ResumeAsset.Kind.AVATAR,
            ).first()
            if avatar and not (avatar.metadata or {}).get('revoked_at') and avatar.size_bytes <= 3_000_000:
                with avatar.file.open('rb') as handle:
                    avatar_bytes = handle.read(3_000_001)
                if len(avatar_bytes) > 3_000_000:
                    avatar_bytes = None
        content, mime_type, page_count = _render_with_rendercv(
            resume_json,
            design_json,
            artifact.format,
            avatar_bytes=avatar_bytes,
        )
        extension = 'pdf' if artifact.format == ResumeArtifact.Format.PDF else 'png'
    else:
        raise RenderFailure('format_unsupported', '不支持的导出格式。')
    checksum = __import__('hashlib').sha256(content).hexdigest()
    version_label = f'v{artifact.content_version.version_number}' if artifact.content_version_id else artifact.draft_etag[:12]
    filename = f'resume-{artifact.resume_id}-{version_label}.{extension}'
    asset = ResumeAsset(
        resume=artifact.resume,
        kind=ResumeAsset.Kind.ARTIFACT,
        original_name=filename,
        mime_type=mime_type,
        size_bytes=len(content),
        checksum_sha256=checksum,
        metadata={'renderer': RENDERER_NAME, 'renderer_version': RENDERER_VERSION},
    )
    asset.file.save(filename, ContentFile(content), save=False)
    asset.save()
    artifact.asset = asset
    artifact.page_count = page_count
    artifact.status = ResumeArtifact.Status.READY
    artifact.completed_at = timezone.now()
    artifact.save(update_fields=['asset', 'page_count', 'status', 'completed_at'])
    return artifact
